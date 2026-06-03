import re
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

base = 'd:/cc 内容/新建文件夹/manuscript'

with open(os.path.join(base, 'sections/introduction_zh.tex'), 'r', encoding='utf-8') as f:
    intro = f.read()
with open(os.path.join(base, 'sections/results_discussion_zh.tex'), 'r', encoding='utf-8') as f:
    rd = f.read()
with open(os.path.join(base, 'sections/methods_zh.tex'), 'r', encoding='utf-8') as f:
    methods = f.read()
with open(os.path.join(base, 'references.bib'), 'r', encoding='utf-8') as f:
    refs = f.read()

def latex_to_text(tex):
    tex = re.sub(r'\\section\{([^}]*)\}', r'\n\n\1\n\n', tex)
    tex = re.sub(r'\\subsection\{([^}]*)\}', r'\n\n\1\n\n', tex)
    tex = re.sub(r'\\cite\{[^}]*\}', '', tex)
    tex = tex.replace('$', '')
    tex = tex.replace('{', '')
    tex = tex.replace('}', '')
    tex = tex.replace('~', ' ')
    tex = tex.replace('\\#', '#')
    tex = tex.replace('\\cdot', '·')
    tex = tex.replace('\\degree', '°')
    tex = tex.replace('\\textsubscript', '')
    tex = tex.replace('\\%', '%')
    tex = re.sub(r'\\[a-zA-Z]+', '', tex)
    tex = re.sub(r'  +', ' ', tex)
    tex = re.sub(r'\n\s*\n', '\n\n', tex)
    return tex.strip()

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(11)

# === TITLE (Chinese) ===
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('具有"三明治"结构的 Bi/MXene 复合负极\n用于同时实现高倍率与高稳定性钠离子存储')
run.bold = True
run.font.size = Pt(15)
run.font.name = 'Times New Roman'

# Authors
author = doc.add_paragraph()
author.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = author.add_run('[作者及单位 — 待补充]')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(150, 150, 150)

# === ABSTRACT ===
doc.add_paragraph()
abs_h = doc.add_paragraph()
run = abs_h.add_run('摘  要')
run.bold = True
run.font.size = Pt(13)
abs_p = doc.add_paragraph('[全文完成后撰写]')
abs_p.paragraph_format.first_line_indent = Inches(0.3)

# === INTRODUCTION ===
doc.add_paragraph()
intro_h = doc.add_paragraph()
run = intro_h.add_run('1. 引言')
run.bold = True
run.font.size = Pt(13)

for p in latex_to_text(intro).split('\n\n'):
    p = p.strip()
    if not p or p == '引言':
        continue
    para = doc.add_paragraph()
    para.paragraph_format.first_line_indent = Inches(0.3)
    para.paragraph_format.space_after = Pt(6)
    para.paragraph_format.line_spacing = 1.5
    run = para.add_run(p)
    run.font.size = Pt(11)

# === R&D ===
doc.add_paragraph()
rd_h = doc.add_paragraph()
run = rd_h.add_run('2. 结果与讨论')
run.bold = True
run.font.size = Pt(13)

current_sub = None
for p in latex_to_text(rd).split('\n\n'):
    p = p.strip()
    if not p:
        continue
    # Check if it's a subsection heading
    known_subs = [
        'Bi/MXene 复合材料的合成与结构表征',
        '电化学储钠性能',
        '反应动力学与电荷存储机理',
        '循环后结构分析',
    ]
    is_heading = False
    for sub in known_subs:
        if sub in p and len(p) < 80:
            is_heading = True
            current_sub = sub
            break
    if p in ('结果与讨论', '结论'):
        continue
    if is_heading:
        sub_p = doc.add_paragraph()
        run = sub_p.add_run(p)
        run.bold = True
        run.font.size = Pt(12)
    else:
        para = doc.add_paragraph()
        para.paragraph_format.first_line_indent = Inches(0.3)
        para.paragraph_format.space_after = Pt(6)
        para.paragraph_format.line_spacing = 1.5
        run = para.add_run(p)
        run.font.size = Pt(11)

# === METHODS ===
doc.add_paragraph()
m_h = doc.add_paragraph()
run = m_h.add_run('3. 实验部分')
run.bold = True
run.font.size = Pt(13)

known_method_subs = [
    'Ti₃C₂Tₓ MXene 的合成',
    'Bi/MXene 复合材料的合成',
    '材料表征',
    '电化学测试',
]

for p in latex_to_text(methods).split('\n\n'):
    p = p.strip()
    if not p:
        continue
    if p == '实验部分':
        continue
    is_heading = False
    for sub in known_method_subs:
        if sub in p and len(p) < 80:
            is_heading = True
            break
    if is_heading:
        sub_p = doc.add_paragraph()
        run = sub_p.add_run(p)
        run.bold = True
        run.font.size = Pt(12)
    else:
        para = doc.add_paragraph()
        para.paragraph_format.first_line_indent = Inches(0.3)
        para.paragraph_format.space_after = Pt(6)
        para.paragraph_format.line_spacing = 1.5
        run = para.add_run(p)
        run.font.size = Pt(11)

# === REFERENCES ===
doc.add_paragraph()
ref_h = doc.add_paragraph()
run = ref_h.add_run('参考文献')
run.bold = True
run.font.size = Pt(13)

for line in refs.strip().split('\n'):
    line = line.strip()
    if line:
        para = doc.add_paragraph()
        para.paragraph_format.space_after = Pt(2)
        para.paragraph_format.line_spacing = 1.15
        run = para.add_run(line)
        run.font.size = Pt(10)

output_path = os.path.join(base, 'Bi-MXene_SIB_paper_draft_zh.docx')
doc.save(output_path)
print(f'Done: {output_path}')
