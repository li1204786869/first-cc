import re
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

base = 'd:/cc 内容/新建文件夹/manuscript'

with open(os.path.join(base, 'sections/introduction.tex'), 'r', encoding='utf-8') as f:
    intro = f.read()
with open(os.path.join(base, 'sections/results_discussion.tex'), 'r', encoding='utf-8') as f:
    rd = f.read()
with open(os.path.join(base, 'sections/methods.tex'), 'r', encoding='utf-8') as f:
    methods = f.read()
with open(os.path.join(base, 'references.bib'), 'r', encoding='utf-8') as f:
    refs = f.read()

def clean_tex(text):
    """Strip LaTeX commands to readable text."""
    # Save math for later restoration
    text = re.sub(r'\$([^$]+)\$', r' \1 ', text)
    # Remove section commands but keep title text
    text = re.sub(r'\\section\{([^}]*)\}', r'\n[SECTION]\n\1\n', text)
    text = re.sub(r'\\subsection\{([^}]*)\}', r'\n[SUBSECTION]\n\1\n', text)
    # Remove cite
    text = re.sub(r'\\cite\{[^}]*\}', '', text)
    # Remove other LaTeX
    text = text.replace('{', '').replace('}', '').replace('~', ' ')
    text = text.replace('\\#', '#')
    text = text.replace('\\cdot', '·')
    text = text.replace('\\degree', '°')
    text = text.replace('\\%', '%')
    text = re.sub(r'\\[a-zA-Z]+', '', text)
    text = re.sub(r'  +', ' ', text)
    return text.strip()

def write_paragraphs(doc, text, section_titles_to_skip=None):
    """Split text by blank lines and add as paragraphs, skipping known headings."""
    if section_titles_to_skip is None:
        section_titles_to_skip = set()

    blocks = [b.strip() for b in re.split(r'\n\s*\n', text) if b.strip()]

    for block in blocks:
        # Check if this is a section/subsection marker
        is_section = block.startswith('[SECTION]')
        is_subsection = block.startswith('[SUBSECTION]')

        if is_section or is_subsection:
            title = block.split('\n', 1)[1].strip() if '\n' in block else block.replace('[SECTION]', '').replace('[SUBSECTION]', '').strip()
            if title in section_titles_to_skip:
                continue
            p = doc.add_paragraph()
            r = p.add_run(title)
            r.bold = True
            r.font.size = Pt(13) if is_section else Pt(12)
            r.font.name = 'Times New Roman'
        else:
            if len(block) < 5:
                continue
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Inches(0.3)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = 1.5
            r = p.add_run(block)
            r.font.size = Pt(11)
            r.font.name = 'Times New Roman'

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(11)

# === TITLE ===
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run('Bi/MXene Composite Anode with Sandwich Architecture for\nSimultaneously High-Rate and Stable Sodium-Ion Storage')
r.bold = True
r.font.size = Pt(15)
r.font.name = 'Times New Roman'

author = doc.add_paragraph()
author.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = author.add_run('[Authors and Affiliations — to be added]')
r.font.size = Pt(10)
r.font.color.rgb = RGBColor(150, 150, 150)
r.font.name = 'Times New Roman'

# === ABSTRACT ===
doc.add_paragraph()
abs_h = doc.add_paragraph()
r = abs_h.add_run('Abstract')
r.bold = True
r.font.size = Pt(13)
r.font.name = 'Times New Roman'

abs_text = (
    "Addressing the challenge of structural collapse caused by volume expansion in alloy-type anodes "
    "for sodium-ion batteries, this study proposes a universal strategy based on constructing a "
    "nanoconfined system using two-dimensional MXene, achieving atomic-level encapsulation of Bi "
    "nanoparticles between MXene interlayers and forming a robust \"sandwich\" confined structure. "
    "Multi-scale characterization confirms that this architecture utilizes the mechanical strength of "
    "MXene for physical confinement, combined with covalent Ti–O–Bi bonds for chemical anchoring, "
    "fundamentally suppressing particle pulverization of Bi over ultra-long cycling. The resulting "
    "composite electrode exhibits exceptional structural stability, withstanding 5,000 cycles at "
    "2 A g⁻¹ without significant capacity decay, and demonstrates outstanding rate response "
    "(up to 20 A g⁻¹). EIS and GITT analyses reveal a strong synergistic effect between the "
    "MXene conductive network and the shortened diffusion pathways arising from Bi nanoscaling, "
    "greatly enhancing charge-transfer kinetics. This work demonstrates the immense potential of "
    "the \"physical confinement–chemical anchoring\" synergistic mechanism in stabilizing "
    "high-capacity alloying electrodes, providing important theoretical guidance for designing "
    "next-generation high-energy-density, long-life sodium-ion batteries."
)

abs_p = doc.add_paragraph()
abs_p.paragraph_format.first_line_indent = Inches(0.3)
abs_p.paragraph_format.line_spacing = 1.5
r = abs_p.add_run(abs_text)
r.font.size = Pt(11)
r.font.name = 'Times New Roman'

# === INTRODUCTION ===
doc.add_paragraph()
intro_text = clean_tex(intro)
write_paragraphs(doc, intro_text, section_titles_to_skip={'Introduction', '引言'})

# === RESULTS & DISCUSSION ===
doc.add_paragraph()
rd_text = clean_tex(rd)
write_paragraphs(doc, rd_text, section_titles_to_skip={'Results and Discussion', '结果与讨论', 'Conclusion', '结论'})

# === METHODS ===
doc.add_paragraph()
methods_text = clean_tex(methods)
write_paragraphs(doc, methods_text, section_titles_to_skip={'Experimental Section', '实验部分'})

# === REFERENCES ===
doc.add_paragraph()
ref_h = doc.add_paragraph()
r = ref_h.add_run('References')
r.bold = True
r.font.size = Pt(13)
r.font.name = 'Times New Roman'

for line in refs.strip().split('\n'):
    line = line.strip()
    if line:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.15
        r = p.add_run(line)
        r.font.size = Pt(10)
        r.font.name = 'Times New Roman'

output_path = os.path.join(base, 'Bi-MXene_SIB_paper_draft.docx')
doc.save(output_path)
print(f'Done: {output_path}')
